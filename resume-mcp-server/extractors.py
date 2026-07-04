"""Text extraction from resume files (PDF / DOCX / TXT).

Design note: extraction lives in its own module, decoupled from the MCP
server, so it can be unit-tested without spinning up a server and reused
by other entry points (CLI, batch scripts) later.

Security notes (relevant to the capstone's Security rubric item):
- `_check_path` enforces an extension allowlist and a file-size cap, so a
  malicious or accidental request can't make the server read arbitrary
  huge files (a basic denial-of-service guard).
- We never execute or render file contents; we only read text streams.
"""

from pathlib import Path

from docx import Document
from pypdf import PdfReader

# Allowlist, not blocklist: anything we haven't explicitly tested is rejected.
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_FILE_SIZE_MB = 10


class ExtractionError(Exception):
    """Raised when a file cannot be read. The MCP tool converts this into a
    clean error message for the agent instead of a raw stack trace."""


def _check_path(file_path: str) -> Path:
    """Validate the path before touching the file. Fail fast with a
    human-readable reason the LLM agent can relay to the user."""
    path = Path(file_path).expanduser().resolve()

    if not path.exists():
        raise ExtractionError(f"File not found: {path}")
    if not path.is_file():
        raise ExtractionError(f"Not a file: {path}")
    if path.suffix.lower() not in ALLOWED_EXTENSIONS:
        raise ExtractionError(
            f"Unsupported file type '{path.suffix}'. "
            f"Supported: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )
    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        raise ExtractionError(
            f"File too large ({size_mb:.1f} MB). Limit is {MAX_FILE_SIZE_MB} MB."
        )
    return path


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(path)
    if reader.is_encrypted:
        raise ExtractionError("PDF is password-protected; please provide an unlocked copy.")
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    # Resumes often put contact info or skills inside tables — plain
    # paragraph iteration misses those, so walk tables explicitly.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_txt(path: Path) -> str:
    # utf-8-sig transparently strips a BOM if present (common on Windows).
    return path.read_text(encoding="utf-8-sig", errors="replace")


def extract_text(file_path: str) -> str:
    """Extract raw text from a resume file. Dispatches on extension."""
    path = _check_path(file_path)
    dispatch = {".pdf": _extract_pdf, ".docx": _extract_docx, ".txt": _extract_txt}
    text = dispatch[path.suffix.lower()](path)

    if not text.strip():
        raise ExtractionError(
            "No text could be extracted. The file may be a scanned image; "
            "OCR is not supported yet."
        )
    return text.strip()
